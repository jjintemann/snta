#!/usr/bin/env python

# required python modules:  python-docx, pandas and xlrd

# Qualtrics excel spreadsheet header cheat sheet
#
#	Header		What it is
#	'1A_1'		Professor First Name
#	'1A_2'		Professor Last Name
#	'1A_3'		Professor Email
#	'1B_1'		Student First Name
#	'1B_2'		Student Last Name
#	'1B_3'		Student Email
#	2		Nomination Letter
#	' '		Did Professor ask student for nomination?
#	' .1'		Does student want to be anonymous?

# NOTE: if the script produces no data, it's probably having a KeyError Exception in read_xls.
#       It looks for EOF by catching a KeyError and thinks it's the end of file instead of a real error.

import os
import sys
import pandas
import docx
import traceback


def quit():
    sys.exit()

def remove_spaces(text):
    while text[0] == ' ':
        text = text[1:]
    while text[-1:] == ' ':
        text = text[:-1]
    while text.find('  ') != -1:
        text = text.replace('  ', ' ')
    return text

def read_xls(file):
    xl = pandas.ExcelFile(file)
    xl_sheet = xl.parse(xl.sheet_names[0])


    data = []
    row = 1

    while True:
        data_entry = {
            'prof_first_name': None,
            'prof_last_name': None,
            'prof_email': None,
            'user_first_name': None,
            'user_last_name': None,
            'user_email': None,
            'letter': None,
            'asked': False,
            'anonymous': False
        }

        row += 1
        try:
            if xl_sheet['1A_1'][row]:
                data_entry['prof_first_name'] = remove_spaces(str(xl_sheet['1A_1'][row]))
            if xl_sheet['1A_2'][row]:
                data_entry['prof_last_name'] = remove_spaces(str(xl_sheet['1A_2'][row]))
            if xl_sheet['1A_3'][row]:
                data_entry['prof_email'] = remove_spaces(str(xl_sheet['1A_3'][row]))
            if xl_sheet['1B_1'][row]:
                data_entry['user_first_name'] = remove_spaces(str(xl_sheet['1B_1'][row]))
            if xl_sheet['1B_2'][row]:
                data_entry['user_last_name'] = remove_spaces(str(xl_sheet['1B_2'][row]))
            if xl_sheet['1B_3'][row]:
                data_entry['user_email'] = remove_spaces(str(xl_sheet['1B_3'][row]))
            if xl_sheet[2][row]:
                data_entry['letter'] = remove_spaces(str(xl_sheet[2][row]))

            if xl_sheet[' '][row] and str(xl_sheet[' '][row]) == 'I was not asked by my instructor to nominate them.':
                    data_entry['asked'] = False
            else:
                data_entry['asked'] = True
            if xl_sheet[' .1'][row] and str(xl_sheet[' .1'][row]) == 'I do not want my name shared.':
                data_entry['anonymous'] = True
            else:
                data_entry['anonymous'] == False
            if (
                data_entry['prof_first_name'] or
                data_entry['prof_last_name'] or
                data_entry['prof_email'] or
                data_entry['user_first_name'] or
                data_entry['user_last_name'] or
                data_entry['user_email'] or
                data_entry['letter']
            ):
                data.append(data_entry)
            else:
                raise IOError
        except IOError:
            print("Found a row with no data! May have reached end of file?")
            if row > 300:
                break
        except KeyError:
            print("Data from qualtrics survey spreadhsheet has been imported. Total entries recorded: %s" % len(data))
            break
        except Exception as e:
            print("Error: %s" % e)
            print(traceback.format_exc())
            break
    return data

def organize_results(data):
    results = []

    for entry in data:
        found_prof = False
        for result in results:
            prof_name = (entry['prof_first_name'] + " " + entry['prof_last_name']).title().replace('  ', ' ')
            if (
                entry['prof_email'].lower() == result['prof_email'].lower() or
                prof_name.lower() == result['prof_name'].lower()
            ):
                found_prof = True
                nomination = {
                    'student_name': (entry['user_first_name'] + " " + entry['user_last_name']).title().replace('  ', ' '),
                    'letter': entry['letter'],
                    'student_email': entry['user_email'],
                    'anonymous': entry['anonymous'],
                    'asked': entry['asked']
                }

                result['nominations'].append(nomination)

        if not found_prof:
            new_result = {
                'prof_email': entry['prof_email'],
                'prof_name': (entry['prof_first_name'] + " " + entry['prof_last_name']).title().replace('  ', ' '),
                'nominations': []
            }

            nomination = {
                'student_name': (entry['user_first_name'] + " " + entry['user_last_name']).title().replace('  ', ' '),
                'letter': entry['letter'],
                'student_email': entry['user_email'],
                'anonymous': entry['anonymous'],
                'asked': entry['asked']
            }

            new_result['nominations'].append(nomination)
            results.append(new_result)
    return results


def output_data_file(info, filename):
    doc = docx.Document()
    doc.add_paragraph("Nominee: %s\nEmail: %s\nNumber of nominations: %s" % (info['prof_name'], info['prof_email'], len(info['nominations'])))
    count = 0
    for nomination in info['nominations']:
        count += 1
        if nomination['asked']:
            doc.add_paragraph("Nomination #%s\nTHIS STUDENT INDICATED THEY WERE ASKED BY THE PROFESSOR TO NOMINATE THEM.\n%s" % (count, nomination['letter']))
        else:
            doc.add_paragraph("Nomination #%s\n%s" % (count, nomination['letter']))
    doc.save(filename)


def dupe_check(info):
    for item in info[:]:
        email_list = []
        for i in item['nominations']:
            if i['student_email'].lower() in email_list:
                print("WARNING: Student nominated the same professor more than once!")
                print("%s (%s) has more than one nomination for %s" % (i['student_name'], i['student_email'], item['prof_name']))
                prompt = ''
                while prompt.lower() != 'y' and prompt.lower() != 'n':
                    input("Delete duplicate nomination from final report? (y/n): ")
                if prompt.lower() == 'y':
                    item['nominations'].remove(i)
            else:
                email_list.append(i['student_email'].lower())
    return info


def get_xls():
    wd = os.path.dirname(os.path.abspath(__file__))
    files = os.listdir(wd)
    xls_files = []
    for file in files:
        if file.lower().endswith('.xlsx'):
            xls_files.append(os.path.join(wd, file))
    if len(xls_files) == 0:
        print("No excel spreadsheet found! You must put an xlsx file into the same directory as this program!")
        prompt = input("Hit enter to quit")
        quit()
    if len(xls_files) == 1:
        return xls_files[0]
    print("More than one excel spreadsheet was found. Please select which file is the file to be used.")
    print()
    count = 0
    for file in xls_files:
        count += 1
        print("%s. %s" % (count, os.path.split(file)[1]))

    print()
    prompt = 0

    while prompt < 1 or prompt > len(xls_files):
        prompt = input("Enter file number: ")
        try:
            if prompt[-1:] == '.':
                prompt = prompt[:-1]
            assert prompt.find('.') == -1
            prompt = int(prompt)
        except ValueError:
            prompt = 0
        except AssertionError:
            prompt = 0

    return xls_files[prompt-1]

def main(file):
    print("Importing data from %s" % file)
    qualtrics_results = read_xls(file)
    organized_data = organize_results(qualtrics_results)
    results = dupe_check(organized_data)

    finalist_report = "Finalists for Student Nominated Teaching Award:\n\n"
    r = []
    for result in results:
        if len(result['nominations']) >= required_nominations:
            output_data_file(result, result['prof_name'] + " (2019 SNTA Finalist).docx")
            finalist_report += "%s (%s nominations)\n" % (result['prof_name'], len(result['nominations']))
        r.append(result['prof_name'] + " (" + str(len(result['nominations'])) + " nominations)")

    print(finalist_report)
    r = list(set(r))
    r.sort()
    print("Summary of Nominations:")
    for i in r:
        print(i)

if __name__ == '__main__':
    print("     University of Green Bay")
    print("Student Nominated Teaching Awards")
    print("   Qualtrics Data Analysis Tool")
    print()
    print()
    print("This application is designed to take SNTA qualtrics data and produce word doc summaries for each nomineee.")
    print("Please make sure the qualtrics xlsx file is in the same folder you are running this app out of.")
    print("Note that this application will only produce summary files for the finalists. The summaries will not contain student names, only their nomination letters, which can be passed along to the finalists.")
    print()
    while True:
        required_nominations = input("How many nominations do you want required for a teacher to be a finalist?: ")
        try:
           required_nominations = int(required_nominations)
           assert required_nominations >= 0
        except Exception:
           print()
           print("Please enter a positive integer.")
           print()
        else:
            break

    file = get_xls()
    main(file)
    print()
    print()
    print("Once you are do reviewing the nomination summary, press Ctrl-C to quit.")
    try:
        while True:
            input("")
    except KeyboardInterrupt:
        quit()

